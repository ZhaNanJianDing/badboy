import re
import datetime
import matplotlib.pyplot as plt
import jieba
from wordcloud import WordCloud
import seaborn as sns
from PIL import Image
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from baogao import generator

#解决中文显示问题
plt.rcParams['font.sans-serif']=['SimHei']
plt.rcParams['axes.unicode_minus'] = False

#计算渣男分值
totalPoint = 0

#过滤的词
exclude={'图片','表情','杨进'}
#鉴定结果图片保存到此列表
listOutPut = []
#词频保存列表
listWordCount = []
#保存语录匹配结果
listSentences = []

#自添词汇
jieba.add_word("渣男")
#渣男名字
znName = ""
#聊天记录
filename = ""


#做Days（天）表格
def get_date(data):
    dates = re.findall(r'\d{4}-\d{2}-\d{2}',data)
    days = [date[-2:] for date in dates]
    plt.plot()#生成图片
    sns.countplot(days)#days就是数据集，countplot统计days每一天的重复数量就是每一天聊天数目
    
    #early middle late
    early_month = days.count('01')+days.count('02')+days.count('03')+days.count('04')+days.count('05')+days.count('06')+days.count('07')+days.count('08')+days.count('09')+days.count('10')
    middle_month = days.count('11')+days.count('12')+days.count('13')+days.count('14')+days.count('15')+days.count('16')+days.count('17')+days.count('18')+days.count('19')+days.count('20')
    late_month = days.count('21')+days.count('22')+days.count('23')+days.count('24')+days.count('25')+days.count('26')+days.count('27')+days.count('28')+days.count('29')+days.count('30')+days.count('31')
    #print(days)
    #late_month += 100
    
    
    plt.title('聊天日期分布') 
    plt.savefig("./cons/pic/date.png")#保存图片
    listOutPut.append('./cons/pic/date.png')#结果图片加入图片列表
    
    
    global totalPoint
    if late_month == max(early_month, middle_month, late_month):
        listOutPut.append("月底的聊天频率最高，或许他月底没钱了才想起你？(渣男指数+10)")
        totalPoint += 10
    else:
        listOutPut.append("看起来他不是月底没钱了才想起你的渣男…(渣男指数+0)")
    plt.show()
    
    
    #isocalendar()ISO标准化日期，返回三个值的元组（年份，星期数week number，星期几weekday）星期一为1，星期天为7
    weekdays = [datetime.date(int(date[:4]),int(date[5:7]),int(date[-2:])).isocalendar()[-1]  for date in dates]
    plt.plot()
    sns.countplot(weekdays)
    
    weekday_count = weekdays.count('1') + weekdays.count('2') + weekdays.count('3') + weekdays.count('4') + weekdays.count('5')
    weekend_count = weekdays.count('6') + weekdays.count('7')
        
    
    plt.title('星期几聊天最多？')
    plt.savefig("./cons/pic/datetime.png")#保存图片
    listOutPut.append('./cons/pic/datetime.png')
    if weekday_count >= weekend_count:
        listOutPut.append("看起来他不是周末寂寞了才想起你的渣男…(渣男指数+0)")
    else:
        totalPoint += 10
        listOutPut.append("似乎你们在周末聊得更多，或许他是寂寞了才想起你？(渣男指数+10)")
    
    
    
    plt.show()
   


#一天24h的图
def get_time(data):
    times = re.findall(r'\d+:\d{2}:\d{2}',data)  
    hours = [time.split(":")[0] for time in times]#对每一个time分割出代表小时的那部分 
    plt.plot()
    #因为qq消息导出格式里，小时0-9是以个位形式出现，所以用个位的来对应
    sns.countplot(hours,order=['6','7','8','9','10','11','12','13','14','15','16','17','18','19','20','21','22','23','0','1','2','3','4','5'])
    
    night_count = hours.count('0') + hours.count('1') + hours.count('2') + hours.count('3') + hours.count('4') + hours.count('5')
    day_count = hours.count('6') + hours.count('7') + hours.count('8') + hours.count('9') + hours.count('10') + hours.count('11') + hours.count('12') + hours.count('13') + hours.count('14') + hours.count('15') + hours.count('16') + hours.count('17') + hours.count('18') + hours.count('19') + hours.count('20') + hours.count('21') + hours.count('22') + hours.count('23')
    
    plt.title("通常在什么时候聊天？")
    plt.savefig("./cons/pic/time.png")#保存图片
    listOutPut.append('./cons/pic/time.png')
    
    global totalPoint
    if night_count >= day_count:
        totalPoint += 20
        listOutPut.append("熬夜有黑眼圈的渣男！(渣男指数+20)")
    else:
        listOutPut.append("也许他没有黑眼圈？(渣男指数+0)")
    plt.show()


#词云制作
def get_wordcloud(text_data):
    # word_list 为一个list对象，其中每一个元素为 '一句 聊天 记录 分词 结果' 的形式
    word_list = [" ".join(jieba.cut(sententce)) for sententce in text_data]
    # new_text 为 str对象，内容为 "所有 的 句子 合在 一起 单词 用 空格 隔开"
    new_text = ' '.join(word_list)

    #背景图片的设置，设置词云形状 
    pic = Image.open('./material/zhanan.png')
    mang_mask = np.array(pic) 
    plt.plot()
    wordcloud = WordCloud(
        background_color="white",
        font_path = '/home/shen/Downloads/fonts/msyh.ttc',
        mask = mang_mask,
        #设置最大词数
        max_words=100,
        #停用词
        stopwords=exclude,).generate(new_text)
    plt.imshow(wordcloud)
    plt.axis('off')
    wordcloud.to_file('./cons/pic/example1.png')
    listOutPut.append('./cons/pic/example1.png')
    plt.show()


#统计词出现的次数    
def count_number(text_data):   
    count_text = ' '.join(text_data)
    words = jieba.cut(count_text)
    
    #统计词频并排序    
    counts = {}
    exclude_flag = 0
    for word in words:        
        for i in exclude: 
            if(word == i):
                exclude_flag = 1        
        if ((len(word) == 1) or (exclude_flag == 1)):
            exclude_flag = 0
            continue
        else:
            counts[word] = counts.get(word, 0) + 1
            
    items = list(counts.items())
    items.sort(key = lambda x:x[1], reverse=True)
    
    #词语和频率存入列表备用
    for i in range(20):
        (word, count) = items[i]
        cipin = str(word)+":"+str(count)+"次"
        listWordCount.append(cipin)

    #将统计出来的关键词存到txt文件备用
    f_keywords = open("./material/keywords.txt", "w")
    for i in range(20):
        (word, count) = items[i]
        f_keywords.write("{0}\n".format(word))
    f_keywords.close()


#匹配文本内容，调用上一个函数
def get_content(data):
    pa = re.compile(r'\d{4}-\d{2}-\d{2}.*?\d+:\d{2}:\d{2}.*?\n(.*?)\n\n',re.DOTALL)  
    content = re.findall(pa,data)  #此处为list对象，将聊天记录中的其他内容拆去，仅保留对话内容，每句对话为list对象中的一个元素
    get_wordcloud(content)
    count_number(content)



def match_sentence(sentences, keywords):
    plot_list = []    # 声明一个空的列表，用来存储所有包含关键词的句子
    contents=[]     # 声明一个空的列表，用来存储所有关键词
    f = open(keywords, 'r')
    while True:
        line=f.readline()
        line=line.strip('\n')
        if line:
            contents.append(line)
        else:
            break
    f.close()
    #print(contents)
    with open(sentences, 'r') as fs:
        st = fs.read()
    fs.close()    
    for kw in contents:
        pattern = re.compile('[，。,.！？!?、；;]*' + kw + '[^，。,.？！]*[，。？*！;；]')
        pattern = re.compile('[^，。,.!?！？、：;]*' + kw + '[^，。,.？！]*[，。？*！;；]')
        plot_temp = re.findall(pattern, st)    # 搜索句子，把结果缓存到plot_temp
        plot_list.append(' '.join(plot_temp))    # 把这次搜索结果添加都结果列
        #print(plot_temp) 

    #print(plot_list)
    while '' in plot_list:
        plot_list.remove('')
    #print(plot_list)        
    output_list = '\n\n'.join(plot_list)
    word_cloud_list = ' '.join(plot_list)
    print(output_list)    # 循环搜索完毕就能输出所有结果了
    
    #词云对象
    w = WordCloud(
        width = 600,
        height = 400,
        background_color="white",
        font_path = '/home/shen/Downloads/fonts/msyh.ttc',
        ).generate(word_cloud_list)

    print("\n\n潜在的渣男发言：")
    plt.imshow(w)
    plt.axis('off')
    w.to_file('./cons/pic/yulu.png')
    plt.show()
    count_sentence = "\n统计结果：\n20个他常用的关键词中能匹配到渣男语句的有" + str(len(plot_list)) + "个词    (渣男指数+"+str(len(plot_list)*3)+")"
    listSentences.append(count_sentence)
    listSentences.append("\n")
    listSentences.append("根据高频词，我们猜测，他的潜在渣男发言：")
    global totalPoint
    totalPoint += (len(plot_list))*3
    
    
def end_word(Point, Name):
    if Point >= 60:
        w = Name + "是铁渣男"
    if Point < 60 and Point>30:
        w = Name + "好像不怎么渣"
    if Point <= 30:
        w = Name + "是绝世好男人"
    content = generator(w, length=500)
    return content

#生成word版的鉴定报告
def creat_word(wordName):
    doc = Document()    # doc对象
    doc.add_paragraph("鉴定报告")   # 添加文字
    
    change_flag = 0
    for outputItem in listOutPut:
        if (change_flag % 2) == 0:
            doc.add_picture(outputItem, width=Inches(5))     # 添加图, 设置宽度
        else:
            doc.add_paragraph(outputItem)
        change_flag += 1
    strWordCount = "高频词语:"+'\n'+'    '.join(listWordCount) + '\n'    
    doc.add_paragraph(strWordCount)
    print(strWordCount)
    doc.add_paragraph(listSentences)
    doc.add_picture("./cons/pic/yulu.png", width=Inches(5))
    
    zuizhongdefen = "被试者的最终渣男指数为：" + str(totalPoint) + "\n" + end_word(totalPoint, znName)
    doc.add_paragraph(zuizhongdefen)
    
   


 

    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    
    print(totalPoint)
    print(end_word(totalPoint, znName))
        
    doc.save(wordName)     # 保存路径



def zhanan():
    global filename
    #encoding一下，防止乱码
    with open(filename,encoding="UTF-8") as f:
        data = f.read()
    get_date(data)
    get_time(data)
    get_content(data)
    sentences = "./material/sentences.txt"
    keywords = "./material/keywords.txt"
    match_sentence(sentences, keywords)
    
    
    
    creat_word('./cons/baogao.docx')
 
if __name__ == '__main__':
    zhanan()
